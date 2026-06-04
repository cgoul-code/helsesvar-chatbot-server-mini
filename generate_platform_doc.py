"""Generate a Word document describing the technical platform of the
HEI chatbot server and client. Run once: `python generate_platform_doc.py`.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


doc = Document()

# Title
title = doc.add_heading("Teknisk plattform – HEI chatbot", level=0)

intro = doc.add_paragraph()
intro.add_run(
    "Kort oversikt over teknologiplattformen som serveren og klienten "
    "(som svarer på spørsmål) er bygget på."
).italic = True

# --- Server ---
doc.add_heading("Server (llama_chatbot_server_mini)", level=1)
doc.add_paragraph("Python-basert, asynkron RAG-/agenttjeneste.")

add_table(
    doc,
    ["Lag", "Teknologi"],
    [
        ["Web-rammeverk", "Quart (async Flask-API) + quart-cors, kjørt på Hypercorn (ASGI)"],
        ["Streaming", "Server-Sent Events (SSE) med keepalive-heartbeat"],
        ["RAG / vektorindeks", "LlamaIndex (llama-index 0.14) med FAISS (faiss-cpu) som vektorlager"],
        ["Agent-/arbeidsflyt", "LangGraph + LangChain"],
        ["LLM (svargenerering)", "Pluggbart via LLM_PROVIDER: Azure OpenAI (standard) eller Anthropic Claude"],
        ["Embeddings", "Azure OpenAI Embeddings"],
        ["Sesjon/cache", "diskcache (samtalehistorikk per session_id)"],
        ["Hosting", "Azure (App Service / Functions – detekteres via miljøvariabler)"],
    ],
)
doc.add_paragraph(
    "Indeksene lastes asynkront i bakgrunnen ved oppstart, og /healthz "
    "rapporterer når de er klare."
)

# --- Client ---
doc.add_heading("Klient (chatbot-client-HEI20-v2)", level=1)
doc.add_paragraph("JavaScript-basert single-page-app.")

add_table(
    doc,
    ["Lag", "Teknologi"],
    [
        ["Rammeverk", "React 18"],
        ["Byggverktøy", "Create React App via CRACO (craco start/build)"],
        ["UI-komponenter", "MUI (Material UI) v5 + react-pro-sidebar, react-select, react-tooltip"],
        ["HTTP / streaming", "axios, samt SSE mot serverens /chat og /examples"],
        ["Markdown/visning", "react-native-markdown-display, remark-breaks, react-native-web"],
    ],
)

# --- Summary ---
doc.add_heading("Kort oppsummert", level=1)
for line in [
    "Server: Python + Quart/Hypercorn, RAG med LlamaIndex/FAISS, agentlogikk i "
    "LangGraph/LangChain, LLM fra Azure OpenAI eller Anthropic Claude – på Azure.",
    "Klient: React 18 (CRA/CRACO) med Material UI, snakker med serveren over REST + SSE.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(line)

doc.add_paragraph(
    "Begge er bygget rundt en klassisk RAG-arkitektur (henter relevante dokumenter "
    "fra en vektorindeks og lar en LLM formulere svaret), med en tydelig separasjon "
    "mellom Python-backend og React-frontend."
)

out = "Teknisk_plattform_HEI_chatbot.docx"
doc.save(out)
print(f"Saved {out}")
